import requests 
import pandas as pd 
from lxml import html
from bs4 import BeautifulSoup
from selenium import webdriver
from selenium.webdriver.chrome.options import Options
from docx import Document
import time


# from  docx  import  Document
# w=Document(r'F:\word practice \ a form. docx')
# table_1=w.tables[0]

# # delete row 
# print(len(table_1.rows))
# row2=table_1.rows[1]
# row2._element.getparent().remove(row2._element)
# print(len(table_1.rows))

# # delete column 
# col=table_1.table.columns[1]
# for  cell in  col.cells:
#     cell._element.getparent().remove(cell._element)
# w.save(r'F:\word practice \ delete the rows or columns of the table. docx')


def remove_row(table, row):
    tbl = table._tbl
    tr = row._tr
    tbl.remove(tr)


data = pd.read_excel(r'../SBST.xlsx')
df = pd.DataFrame(data, columns= ['Bus Depot','OBU Serial No'])
dflist = list(df.index)


testDict = {}
busDepotList = []
serialNumList = []
for i in dflist:
    busDepot = (df['Bus Depot'][i]) 
    serialNum = (df['OBU Serial No'][i])
    busDepotList.append(busDepot)
    serialNumList.append(serialNum)
   # testDict[busDepot] = serialNum

#print (busDepotList)
print(serialNumList)

document = Document('test.docx')
tables = document.tables
#ID = 1, bus depot = 5 


# row_count = len(tables[0].rows)
# for i in range(row_count):
#     print("here")
#     for n in range(len(busDepotList)):
#         if str(tables[0].cell(i,1).text) == str(serialNumList[n]):
#             print("here3")
#             tables[0].cell(i,5).text = busDepotList[n]

# row_count = len(tables[0].rows)
# print(row_count)
# for i in range(row_count):
#     print(i)
#     if i == row_count-2:
#         break
#     if str(tables[0].cell(i,0).text) == "":
        # row = tables[0].rows[i]
        # remove_row(tables[0], row)
row_count = len(tables[0].rows)
print(row_count)
for i in range(row_count):
    
    row = tables[0].rows[0]
    #print(i)
    # for n in range(len(busDepotList)):
    if str(tables[0].cell(i,1).text) == str(1):
       # delete row 
        if i == tables[0].rows:
            row2=tables[0].rows[i-1]
            row2._element.getparent().remove(row2._element)
            break
        print(len(tables[0].rows))
        row2=tables[0].rows[i]
        row2._element.getparent().remove(row2._element)
        print(len(tables[0].rows))
        print("here")



document.save('test1.docx')



