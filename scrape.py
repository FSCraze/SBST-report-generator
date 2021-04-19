import requests 
import pandas as pd 
from lxml import html
from bs4 import BeautifulSoup
from selenium import webdriver
from selenium.webdriver.chrome.options import Options
from docx import Document
import time



chrome_options = Options() 
chrome_options.add_experimental_option("detach", True)
driver = webdriver.Chrome(options=chrome_options)


def login():
    driver.get("https://noc-sg-prd.veniam.com")
    driver.find_element_by_id("username").send_keys("starhub_monitoring")
    driver.find_element_by_id("password").send_keys("ry5XRAUU5b0QOzyrYHa95jYr65zHJiZp")
    driver.find_element_by_class_name("mdl-button__ripple-container").click()


def scrape():
    #route to device list 
    driver.get("https://noc-sg-prd.veniam.com/devices?client=bukit_merah&device=0&ticket=0&status=0,1,2,3,4,5&columns=id,ip,serial,version,events,plate,last_ping,last_gps,status,board_uptime,location")
    time.sleep(5)
    
    #click on last ping 
    driver.find_element_by_xpath("//*[@id=\"devices-table-container\"]/table/div/div/ngx-datatable/div/datatable-header/div/div[2]/datatable-header-cell[6]").click()
    time.sleep(5)
    
    #change table view to 200 
    driver.find_element_by_xpath("/html/body/noc-ops/div/devices/div/table-config/div/div[2]/div[2]/mat-form-field/div/div[1]/div").click()
    time.sleep(1)
    driver.find_element_by_xpath("//*[@id=\"mat-option-12\"]").click()
    time.sleep(5)
    
    print("Opening template document")
    #open up word document 
    file = 'format.docx'
    document = Document(file)
    doc_row = 0
    print("Looping through tables on web server")
    #Loop through entire table to check for entries with last ping > 1 day or N/A
    for row in range(150):
        temp = driver.find_element_by_xpath("//*[@id=\"devices-table-container\"]/table/div/div/ngx-datatable/div/datatable-body/datatable-selection/datatable-scroller/datatable-row-wrapper[" +str(row+1)+"]/datatable-body-row/div[2]/datatable-body-cell[6]").text
        if (temp.find("d")!=-1 or temp.find("N")!=-1):
            doc_row+=1
            #print(temp)
   # print("doc_row: " + str(doc_row))

    #access table
    tables = document.tables
    table = tables[0]
    indexes = [1,3,6,9] #columns from the website table that is important
    print("Adding data into template document file")
    #loop through entries with last ping > 1 day or N/A and add to table in word doc
    for rowD in range(doc_row):
        for col in range(len(indexes)):
            data = driver.find_element_by_xpath("//*[@id=\"devices-table-container\"]/table/div/div/ngx-datatable/div/datatable-body/datatable-selection/datatable-scroller/datatable-row-wrapper[" +str(rowD+1)+"]/datatable-body-row/div[2]/datatable-body-cell["+str(indexes[col])+ "]").text
            #print(data)
            cell = table.cell(rowD+1, col+1)
            cell.text = data
   

    document.save('test.docx')
    print("Scraping from website successful")


def addBusDepot():
    print("Opening sbst xlsx..")
    data = pd.read_excel(r'../SBST.xlsx')
    df = pd.DataFrame(data, columns= ['Bus Depot','OBU Serial No'])
    dflist = list(df.index)

    print("extracting bus depot + OBU serial # info")
    busDepotList = []
    serialNumList = []
    for i in dflist:
        busDepot = (df['Bus Depot'][i]) 
        serialNum = (df['OBU Serial No'][i])
        busDepotList.append(busDepot)
        serialNumList.append(serialNum)


    document = Document('test.docx')
    tables = document.tables
    #ID = 1, bus depot = 5 

    print("Inserting bus depot info...")
    row_count = len(tables[0].rows)
    for i in range(row_count):
        print("...")
        for n in range(len(busDepotList)):
            if str(tables[0].cell(i,1).text) == "":
                break
            if str(tables[0].cell(i,1).text) == str(serialNumList[n]):
                tables[0].cell(i,5).text = busDepotList[n]
                print(str(i+1) + "/" + str(row_count-1) + "...")
    
    document.save('test.docx')
    print("successful")



login()
scrape()
addBusDepot()







