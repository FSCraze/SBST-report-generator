from docx import Document
document = Document()
font = document.styles['Normal'].font

font.name = 'Calibri'
paragraph = document.add_paragraph('Proactive Monitoring Weekly Report')
paragraph = document.add_paragraph('Date Period 15 - 21 March 2021')
table = document.add_table(rows=2, cols=2)
cell = table.cell(0, 1)
cell.text = 'parrot, possibly dead'
styles = document.styles
table.style = document.styles['Medium Grid 1 Accent 1']
document.save('test2.docx')

# import docx
# from docx.enum.table import WD_TABLE_DIRECTION
# from docx.shared import Cm, Inches

# file = 'test.docx'
# doc = docx.Document(file)
# tbls = doc.tables 
# test = tbls[1]
# test.table_direction = WD_TABLE_DIRECTION.LTR
# test.add_column(Inches(1.0))
# doc.save(file)

def delete_columns(table, columns):
    # sort columns descending
    columns.sort(reverse=True)
    
    grid = table._tbl.find("w:tblGrid", table._tbl.nsmap)
    for ci in columns:
        for cell in table.column_cells(ci):
            cell._tc.getparent().remove(cell._tc)

        # Delete column reference
        col_elem = grid[ci]
        grid.remove(col_elem)

